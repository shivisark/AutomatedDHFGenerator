# DHF and Excel 

import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Path
base_dir = os.path.dirname(__file__)
output_dir = os.path.join(base_dir, "output_reports")
os.makedirs(output_dir, exist_ok=True)

# Input file paths
user_needs_file = os.path.join(base_dir, "user_needs.csv")
design_inputs_file = os.path.join(base_dir, "design_inputs.csv")
design_outputs_file = os.path.join(base_dir, "design_outputs.csv")
verification_methods_file = os.path.join(base_dir, "verification_methods.csv")
change_log_file = os.path.join(base_dir, "change_log.csv")
traceability_matrix_output = os.path.join(output_dir, "traceability_matrix.csv")
risk_table_file = os.path.join(base_dir, "risk_table.csv")
device_overview_file = os.path.join(base_dir, "device_overview.txt")

# Excel output path
excel_output_file = os.path.join(output_dir, "DesignControls.xlsx")

# Combined Excel with the design controls and risk management
with pd.ExcelWriter(excel_output_file, engine='xlsxwriter') as writer:
    def write_sheet(file_path, sheet_name):
        if os.path.exists(file_path):
            df = pd.read_csv(file_path)
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            worksheet = writer.sheets[sheet_name]
            workbook = writer.book

            # Wrap text
            wrap_format = workbook.add_format({'text_wrap': True})
            for col_num in range(len(df.columns)):
                worksheet.set_column(col_num, col_num, 25, wrap_format)

            # Header format
            header_format = workbook.add_format({
                'bold': True, 'text_wrap': True,
                'valign': 'top', 'fg_color': '#DCE6F1', 'border': 1
            })
            for col_num, _ in enumerate(df.columns):
                worksheet.write(0, col_num, df.columns[col_num], header_format)

    # All Input files
    write_sheet(user_needs_file, 'User Needs')
    write_sheet(design_inputs_file, 'Design Inputs')
    write_sheet(design_outputs_file, 'Design Outputs')
    write_sheet(verification_methods_file, 'Verification Methods')
    write_sheet(change_log_file, 'Change Log')
    write_sheet(traceability_matrix_output, 'Traceability Matrix')

    # RPN Calculation
    if os.path.exists(risk_table_file):
        risk_df = pd.read_csv(risk_table_file)
        if all(col in risk_df.columns for col in ['Severity', 'Occurrence', 'Detection']):
            risk_df['RPN'] = (
                pd.to_numeric(risk_df['Severity'], errors='coerce') *
                pd.to_numeric(risk_df['Occurrence'], errors='coerce') *
                pd.to_numeric(risk_df['Detection'], errors='coerce')
            )
        sheet_name = 'Risk Management'
        risk_df.to_excel(writer, sheet_name=sheet_name, index=False)
        worksheet = writer.sheets[sheet_name]
        wrap_format = writer.book.add_format({'text_wrap': True})
        for col_num in range(len(risk_df.columns)):
            worksheet.set_column(col_num, col_num, 25, wrap_format)

# Word Document

excel_reader = pd.ExcelFile(excel_output_file)
from docx.shared import Inches

doc = Document()

# Set 1-inch margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


# Title page and formating
doc.add_heading('Design History File (DHF)', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
run = p.add_run("Wearable ECG Patch")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph("Automated DHF Project - Shivangi Sarkar", style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_page_break()

# Device_overview.txt inserted
if os.path.exists(device_overview_file):
    doc.add_heading('Device Overview', level=1)
    with open(device_overview_file, 'r') as f:
        overview_text = f.read()
    for line in overview_text.strip().split('\n'):
        doc.add_paragraph(line.strip())
    doc.add_page_break()

# Excel sheets combined to the word document
for sheet_name in excel_reader.sheet_names:
    df = excel_reader.parse(sheet_name)
    doc.add_heading(sheet_name, level=1)

    if df.empty:
        doc.add_paragraph("No data available.")
        continue

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    doc.add_paragraph()
    doc.add_page_break()


# Generating DHF word document
word_output_file = os.path.join(output_dir, "DHF_Report.docx")
doc.save(word_output_file)
